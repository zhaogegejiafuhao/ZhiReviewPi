"""
希沃智教π 学情报告多格式导出服务

支持四种格式：
1. JSON — 原始数据导出
2. CSV — 表格数据导出
3. Word — 学情档案（含雷达图描述+薄弱点列表+建议）
4. PDF — 家长精简版（仅含雷达图描述+前3薄弱点+家庭辅导建议）
"""
import csv
import io
import json
from datetime import date
from typing import Optional


class ExportService:
    """学情报告导出服务"""

    @staticmethod
    def export_json(report_data: dict) -> str:
        """导出JSON格式"""
        return json.dumps(report_data, ensure_ascii=False, indent=2)

    @staticmethod
    def export_csv(report_data: dict) -> str:
        """导出CSV格式（薄弱知识点表格）"""
        output = io.StringIO()
        writer = csv.writer(output)

        # 表头
        writer.writerow(["知识点ID", "知识点名称", "薄弱度", "错题数", "主要错因", "改进建议"])

        # 数据行
        for wp in report_data.get("weak_points", []):
            error_causes = wp.get("error_cause_distribution", {})
            top_cause = max(error_causes, key=error_causes.get) if error_causes else "无"

            writer.writerow([
                wp.get("knowledge_id", ""),
                wp.get("knowledge_name", ""),
                f"{wp.get('weakness_score', 0):.1%}",
                wp.get("error_count", 0),
                top_cause,
                wp.get("suggestion", ""),
            ])

        return output.getvalue()

    @staticmethod
    def export_word(report_data: dict) -> bytes:
        """导出Word学情档案"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            # python-docx未安装，降级为纯文本
            return ExportService._export_text_report(report_data).encode("utf-8")

        doc = Document()

        # 标题
        title = doc.add_heading("学情分析报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本信息
        doc.add_paragraph(f"学生ID: {report_data.get('student_id', 'N/A')}")
        doc.add_paragraph(f"分析日期: {report_data.get('analysis_date', date.today().isoformat())}")
        doc.add_paragraph("")

        # 雷达图数据
        doc.add_heading("知识维度掌握度", level=1)
        radar = report_data.get("radar", {})
        if radar:
            table = doc.add_table(rows=1, cols=2, style="Table Grid")
            table.rows[0].cells[0].text = "知识维度"
            table.rows[0].cells[1].text = "掌握度"
            for dim, score in radar.items():
                row = table.add_row()
                row.cells[0].text = dim
                row.cells[1].text = f"{score:.0%}"

        doc.add_paragraph("")

        # 薄弱知识点
        doc.add_heading("薄弱知识点详情", level=1)
        for wp in report_data.get("weak_points", []):
            doc.add_heading(f"{wp.get('knowledge_name', '未知')}（薄弱度 {wp.get('weakness_score', 0):.0%}）", level=2)

            p = doc.add_paragraph()
            p.add_run(f"错题数量: ").bold = True
            p.add_run(f"{wp.get('error_count', 0)} 道")

            # 错因分布
            error_causes = wp.get("error_cause_distribution", {})
            if error_causes:
                p = doc.add_paragraph()
                p.add_run("错因分布: ").bold = True
                cause_str = "、".join([f"{k}({v}道)" for k, v in error_causes.items()])
                p.add_run(cause_str)

            # 根源分析
            root_cause = wp.get("root_cause")
            if root_cause:
                p = doc.add_paragraph()
                p.add_run("根源分析: ").bold = True
                p.add_run(f"源于{root_cause.get('root_name', '未知')}，贡献率{root_cause.get('contribution_ratio', 0):.0%}")

            # 建议
            p = doc.add_paragraph()
            p.add_run("改进建议: ").bold = True
            p.add_run(wp.get("suggestion", "暂无"))

        # 订正状态
        doc.add_heading("订正状态", level=1)
        correction_status = report_data.get("correction_status", {})
        if correction_status:
            p = doc.add_paragraph()
            p.add_run(f"总错题: {correction_status.get('total_errors', 0)} 道, ")
            p.add_run(f"已订正: {correction_status.get('corrected', 0)} 道, ")
            p.add_run(f"订正率: {correction_status.get('correction_rate', 0):.0%}")

        # 保存到bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def export_pdf(report_data: dict) -> bytes:
        """导出家长精简PDF"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import cm
            from reportlab.pdfgen import canvas
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            # reportlab未安装，降级为纯文本
            return ExportService._export_text_report(report_data).encode("utf-8")

        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        # 尝试注册中文字体
        try:
            pdfmetrics.registerFont(TTFont('SimHei', 'C:/Windows/Fonts/simhei.ttf'))
            font_name = 'SimHei'
        except Exception:
            font_name = 'Helvetica'  # fallback

        y = height - 2 * cm

        # 标题
        c.setFont(font_name, 18)
        c.drawCentredString(width / 2, y, "学情简报（家长版）")
        y -= 1.5 * cm

        c.setFont(font_name, 10)
        c.drawString(2 * cm, y, f"学生: {report_data.get('student_id', 'N/A')}    日期: {report_data.get('analysis_date', '')}")
        y -= 1.5 * cm

        # 雷达图数据
        c.setFont(font_name, 14)
        c.drawString(2 * cm, y, "知识维度掌握度")
        y -= 1 * cm

        c.setFont(font_name, 10)
        radar = report_data.get("radar", {})
        for dim, score in radar.items():
            bar_width = score * 8 * cm
            c.drawString(2 * cm, y, f"{dim}: {score:.0%}")
            y -= 0.6 * cm

        y -= 0.5 * cm

        # 前3薄弱点
        c.setFont(font_name, 14)
        c.drawString(2 * cm, y, "最需关注的薄弱点")
        y -= 1 * cm

        c.setFont(font_name, 10)
        weak_points = report_data.get("weak_points", [])[:3]
        for i, wp in enumerate(weak_points, 1):
            c.drawString(2 * cm, y, f"{i}. {wp.get('knowledge_name', '')}（薄弱度{wp.get('weakness_score', 0):.0%}）")
            y -= 0.6 * cm
            c.drawString(3 * cm, y, f"建议: {wp.get('suggestion', '')}")
            y -= 0.8 * cm

        # 家庭辅导建议
        y -= 0.5 * cm
        c.setFont(font_name, 14)
        c.drawString(2 * cm, y, "家庭辅导建议")
        y -= 1 * cm

        c.setFont(font_name, 10)
        for wp in weak_points[:2]:
            error_causes = wp.get("error_cause_distribution", {})
            top_cause = max(error_causes, key=error_causes.get) if error_causes else ""
            if top_cause:
                advice_map = {
                    "计算粗心": "建议每天安排10分钟口算训练",
                    "概念混淆": "建议陪孩子重新理解核心概念，画思维导图区分易混点",
                    "审题不清": "建议培养读题划重点的习惯",
                    "逻辑跳步": "建议要求孩子每步都写出推理过程",
                }
                advice = advice_map.get(top_cause, "建议关注孩子日常作业完成质量")
                c.drawString(2 * cm, y, f"· {wp.get('knowledge_name', '')}: {advice}")
                y -= 0.6 * cm

        c.showPage()
        c.save()
        return buffer.getvalue()

    @staticmethod
    def _export_text_report(report_data: dict) -> str:
        """纯文本降级格式（当python-docx/reportlab不可用时）"""
        lines = []
        lines.append("=" * 40)
        lines.append("学情分析报告")
        lines.append("=" * 40)
        lines.append(f"学生: {report_data.get('student_id', 'N/A')}")
        lines.append(f"日期: {report_data.get('analysis_date', '')}")
        lines.append("")

        lines.append("【知识维度掌握度】")
        for dim, score in report_data.get("radar", {}).items():
            lines.append(f"  {dim}: {score:.0%}")
        lines.append("")

        lines.append("【薄弱知识点】")
        for wp in report_data.get("weak_points", []):
            lines.append(f"  · {wp.get('knowledge_name', '')} (薄弱度{wp.get('weakness_score', 0):.0%})")
            if wp.get("suggestion"):
                lines.append(f"    建议: {wp['suggestion']}")

        return "\n".join(lines)
